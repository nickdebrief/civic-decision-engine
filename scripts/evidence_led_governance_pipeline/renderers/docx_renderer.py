"""DOCX renderer for the Evidence-Led Governance semantic model."""

from __future__ import annotations

from datetime import date
from pathlib import Path
from typing import Sequence

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from model import (
    Book,
    BulletItem,
    BulletList,
    Callout,
    Chapter,
    CrossReference,
    FlowDiagram,
    FrontMatter,
    PageBreak,
    Paragraph,
    Section,
    Volume,
)
from themes.base import EffectiveTheme, Theme
from themes.handbook import HANDBOOK_THEME
from themes.registry import PUBLICATION_PROFILES


DEFAULT_EFFECTIVE_THEME = EffectiveTheme(
    theme=HANDBOOK_THEME,
    publication_profile=PUBLICATION_PROFILES["digital"],
    page=HANDBOOK_THEME.page,
    title_page=HANDBOOK_THEME.title_page,
    volume_page=HANDBOOK_THEME.volume_page,
    chapter_opening=HANDBOOK_THEME.chapter_opening,
)


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str, size: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for side in ("top", "left", "bottom", "right"):
        node = borders.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_cell_margins(
    cell,
    *,
    top: int,
    bottom: int,
    start: int,
    end: int,
) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def keep_with_next(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    if p_pr.find(qn("w:keepNext")) is None:
        p_pr.append(OxmlElement("w:keepNext"))


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])


class DocxRenderer:
    def __init__(self, theme: EffectiveTheme | Theme | None = None) -> None:
        if theme is None:
            theme = DEFAULT_EFFECTIVE_THEME
        elif isinstance(theme, Theme):
            theme = EffectiveTheme(
                theme=theme,
                publication_profile=PUBLICATION_PROFILES["digital"],
                page=theme.page,
                title_page=theme.title_page,
                volume_page=theme.volume_page,
                chapter_opening=theme.chapter_opening,
            )
        self.effective = theme
        self.theme = theme.theme
        self._bookmark_id = 1

    def render(self, book: Book, out_path: Path) -> Path:
        doc = Document()
        self.setup_styles(doc)
        self.add_header_footer(doc, book.running_title, book.tagline)
        self.title_page(doc, book)

        for block in book.blocks:
            self.render_block(doc, block)

        props = doc.core_properties
        props.title = book.title
        props.subject = book.metadata.get("subject", book.subtitle)
        props.author = book.author
        props.category = book.metadata.get("edition", "")
        props.keywords = book.metadata.get("keywords", "")
        comments = book.metadata.get("comments", "")
        build_identifier = book.metadata.get("build_identifier", "")
        generated_comment = f"Generated manuscript version {book.version}"
        props.comments = " · ".join(part for part in (generated_comment, comments, build_identifier) if part)
        language = book.metadata.get("language", "")
        if language:
            props.language = language

        out_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(out_path)
        Document(out_path)
        return out_path

    def setup_styles(self, doc) -> None:
        theme = self.theme
        normal = doc.styles["Normal"]
        typography = theme.typography
        normal.font.name = typography.body_font
        normal.font.size = Pt(typography.body_size_pt)
        normal.font.color.rgb = rgb(theme.colours.body_text)
        normal.paragraph_format.space_after = Pt(typography.body_space_after_pt)
        normal.paragraph_format.line_spacing = typography.line_spacing
        normal.paragraph_format.widow_control = True

        heading_specs = {
            "Heading 1": theme.headings.heading1,
            "Heading 2": theme.headings.heading2,
            "Heading 3": theme.headings.heading3,
        }
        for style_name, spec in heading_specs.items():
            style = doc.styles[style_name]
            style.font.name = typography.display_font
            style.font.size = Pt(spec.size_pt)
            style.font.bold = spec.bold
            style.font.italic = spec.italic
            style.font.color.rgb = rgb(spec.colour)
            style.paragraph_format.space_before = Pt(spec.space_before_pt)
            style.paragraph_format.space_after = Pt(spec.space_after_pt)
            style.paragraph_format.keep_with_next = True
            style.paragraph_format.widow_control = True
        bullet = doc.styles["List Bullet"]
        bullet.font.name = typography.body_font
        bullet.font.size = Pt(typography.body_size_pt)
        bullet.paragraph_format.left_indent = Inches(theme.headings.bullet_left_indent_inches)
        bullet.paragraph_format.first_line_indent = Inches(-theme.headings.bullet_hanging_indent_inches)

        section = doc.sections[0]
        page = self.effective.page
        section.page_width = Inches(page.width_inches)
        section.page_height = Inches(page.height_inches)
        section.left_margin = Inches(page.margin_left_inches)
        section.right_margin = Inches(page.margin_right_inches)
        section.top_margin = Inches(page.margin_top_inches)
        section.bottom_margin = Inches(page.margin_bottom_inches)
        section.header_distance = Inches(page.header_distance_inches)
        section.footer_distance = Inches(page.footer_distance_inches)
        section.different_first_page_header_footer = (
            theme.header_footer.suppress_first_page_header or theme.header_footer.suppress_first_page_footer
        )

    def add_header_footer(self, doc, running_title: str, tagline: str | None = None) -> None:
        theme = self.theme
        for section in doc.sections:
            hp = section.header.paragraphs[0]
            hp.alignment = self.alignment(theme.header_footer.header_alignment)
            header_text = running_title or theme.header_footer.header_label
            hr = hp.add_run(header_text)
            hr.font.name = theme.typography.body_font
            hr.font.size = Pt(theme.typography.footer_size_pt)
            hr.font.color.rgb = rgb(theme.header_footer.colour)

            fp = section.footer.paragraphs[0]
            fp.alignment = self.alignment(theme.header_footer.footer_alignment)
            footer_text = tagline or theme.header_footer.footer_label
            fr = fp.add_run(f"{footer_text}  ·  ")
            fr.font.name = theme.typography.body_font
            fr.font.size = Pt(theme.typography.footer_size_pt)
            fr.font.color.rgb = rgb(theme.header_footer.colour)
            if theme.header_footer.show_page_number:
                add_page_number(fp)

    @staticmethod
    def alignment(value: str):
        return {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
        }[value]

    def para(
        self,
        doc,
        text: str,
        *,
        align=None,
        italic=False,
        bold=False,
        size=None,
        color=None,
        space_after=None,
        style=None,
        bookmark=None,
    ):
        theme = self.theme
        p = doc.add_paragraph(style=style)
        if align is not None:
            p.alignment = align
        run = p.add_run(text)
        run.font.name = theme.typography.body_font
        run.italic = italic
        run.bold = bold
        if size is not None:
            run.font.size = Pt(size)
        if color is not None:
            run.font.color.rgb = rgb(color) if isinstance(color, str) else color
        if space_after is not None:
            p.paragraph_format.space_after = Pt(space_after)
        if bookmark:
            self.add_bookmark(p, bookmark)
        return p

    def add_bookmark(self, paragraph, name: str) -> None:
        start = OxmlElement("w:bookmarkStart")
        start.set(qn("w:id"), str(self._bookmark_id))
        start.set(qn("w:name"), name)
        end = OxmlElement("w:bookmarkEnd")
        end.set(qn("w:id"), str(self._bookmark_id))
        self._bookmark_id += 1
        paragraph._p.insert(0, start)
        paragraph._p.append(end)

    def add_internal_hyperlink(self, paragraph, text: str, anchor: str):
        theme = self.theme
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("w:anchor"), anchor)
        run = OxmlElement("w:r")
        run_pr = OxmlElement("w:rPr")
        colour = OxmlElement("w:color")
        colour.set(qn("w:val"), self.theme.colours.hyperlink)
        run_pr.append(colour)
        if self.effective.publication_profile.visible_hyperlink_style:
            underline = OxmlElement("w:u")
            underline.set(qn("w:val"), "single")
            run_pr.append(underline)
        text_node = OxmlElement("w:t")
        text_node.text = text
        run.append(run_pr)
        run.append(text_node)
        hyperlink.append(run)
        paragraph._p.append(hyperlink)
        return hyperlink

    def title_page(self, doc, book: Book) -> None:
        theme = self.effective.title_page
        typography = self.theme.typography
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(theme.space_before_pt)
        run = p.add_run(book.title)
        run.font.name = typography.display_font
        run.bold = True
        run.font.size = Pt(theme.title_size_pt)
        run.font.color.rgb = rgb(theme.title_colour)
        p.paragraph_format.space_after = Pt(theme.title_space_after_pt)

        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(book.subtitle)
        r2.font.name = typography.display_font
        r2.italic = True
        r2.font.size = Pt(theme.subtitle_size_pt)
        r2.font.color.rgb = rgb(theme.subtitle_colour)
        p2.paragraph_format.space_after = Pt(theme.subtitle_space_after_pt)

        if theme.show_author:
            self.para(doc, book.author, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=theme.author_size_pt, color=theme.title_colour, space_after=theme.author_space_after_pt)
        if theme.show_version:
            self.para(doc, f"Version {book.version}", align=WD_ALIGN_PARAGRAPH.CENTER, size=theme.metadata_size_pt, color=theme.metadata_colour, space_after=theme.metadata_space_after_pt)
        if theme.show_tagline and book.tagline:
            self.para(doc, book.tagline, align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=theme.metadata_size_pt, color=theme.metadata_colour, space_after=theme.trailing_space_after_pt)
        if theme.show_date:
            self.para(doc, date.today().isoformat(), align=WD_ALIGN_PARAGRAPH.CENTER, size=theme.metadata_size_pt, color=theme.metadata_colour, space_after=theme.metadata_space_after_pt)
        doc.add_page_break()

    def render_block(self, doc, block) -> None:
        if isinstance(block, Volume):
            eyebrow = f"VOLUME {block.number}".strip() if block.number else block.title
            self.part_title_page(doc, eyebrow, block.title, bookmark=block.bookmark)
            for child in block.blocks:
                self.render_block(doc, child)
        elif isinstance(block, FrontMatter):
            heading = doc.add_heading(block.title, level=1)
            if block.bookmark:
                self.add_bookmark(heading, block.bookmark)
            keep_with_next(heading)
            for child in block.blocks:
                self.render_block(doc, child)
        elif isinstance(block, Chapter):
            heading_text = f"Chapter {block.number} — {block.title}" if block.number else block.title
            heading = doc.add_heading(heading_text, level=1)
            heading.paragraph_format.page_break_before = self.effective.chapter_opening.page_break_before
            heading.paragraph_format.space_before = Pt(self.effective.chapter_opening.space_before_pt)
            heading.paragraph_format.space_after = Pt(self.effective.chapter_opening.space_after_pt)
            if block.bookmark:
                self.add_bookmark(heading, block.bookmark)
            keep_with_next(heading)
            for child in block.blocks:
                self.render_block(doc, child)
        elif isinstance(block, Section):
            heading = doc.add_heading(block.heading_text, level=min(max(block.level, 2), 3))
            if block.generated:
                generated = self.theme.generated_sections
                heading.runs[0].font.size = Pt(generated.heading_size_pt)
                heading.runs[0].font.color.rgb = rgb(generated.heading_colour)
                heading.paragraph_format.page_break_before = generated.page_break_before
            if block.bookmark:
                self.add_bookmark(heading, block.bookmark)
            keep_with_next(heading)
            for child in block.blocks:
                self.render_block(doc, child)
        elif isinstance(block, Paragraph):
            self.render_paragraph(doc, block)
        elif isinstance(block, BulletList):
            for item in block.items:
                text = item.text if isinstance(item, BulletItem) else str(item)
                self.para(doc, text, style="List Bullet")
        elif isinstance(block, FlowDiagram):
            self.flow_diagram(doc, block)
        elif isinstance(block, PageBreak):
            doc.add_page_break()
        elif isinstance(block, Callout):
            self.callout_box(doc, block, bookmark=block.bookmark)
        else:
            raise TypeError(f"Unsupported block: {type(block)!r}")

    def render_paragraph(self, doc, paragraph: Paragraph) -> None:
        theme = self.theme
        if paragraph.role == "pagebreak":
            doc.add_page_break()
        elif paragraph.role == "emphasis":
            if paragraph.inline_content:
                self.render_inline_paragraph(doc, paragraph, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, color=theme.headings.emphasis.colour, size=theme.typography.emphasis_size_pt, space_after=theme.headings.emphasis.space_after_pt)
            else:
                self.para(doc, paragraph.text, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, color=theme.headings.emphasis.colour, size=theme.typography.emphasis_size_pt, space_after=theme.headings.emphasis.space_after_pt, bookmark=paragraph.bookmark)
        elif paragraph.role == "bold":
            if paragraph.inline_content:
                self.render_inline_paragraph(doc, paragraph, bold=True, color=theme.headings.body_role_colour)
            else:
                self.para(doc, paragraph.text, bold=True, color=theme.headings.body_role_colour, bookmark=paragraph.bookmark)
        else:
            if paragraph.inline_content:
                self.render_inline_paragraph(doc, paragraph)
            else:
                self.para(doc, paragraph.text, bookmark=paragraph.bookmark)

    def render_inline_paragraph(self, doc, paragraph: Paragraph, **kwargs) -> None:
        p = doc.add_paragraph(style=kwargs.get("style"))
        if kwargs.get("align") is not None:
            p.alignment = kwargs["align"]
        for item in paragraph.inline_content:
            if isinstance(item, CrossReference) and item.target_bookmark:
                self.add_internal_hyperlink(p, item.render_label, item.target_bookmark)
                continue
            text = item.render_label if isinstance(item, CrossReference) else str(item)
            run = p.add_run(text)
            run.font.name = self.theme.typography.body_font
            if kwargs.get("bold"):
                run.bold = True
            if kwargs.get("italic"):
                run.italic = True
            if kwargs.get("color") is not None:
                run.font.color.rgb = rgb(kwargs["color"])
            if kwargs.get("size") is not None:
                run.font.size = Pt(kwargs["size"])
        if kwargs.get("space_after") is not None:
            p.paragraph_format.space_after = Pt(kwargs["space_after"])
        if paragraph.bookmark:
            self.add_bookmark(p, paragraph.bookmark)

    def callout_box(self, doc, callout: Callout, *, bookmark: str | None = None) -> None:
        theme = self.theme
        callout_theme = theme.callouts
        style = callout_theme.styles.get(callout.callout_type, callout_theme.styles["Callout"])
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True

        cell = table.rows[0].cells[0]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, style.fill)
        set_cell_border(cell, style.border, callout_theme.border_size_eighth_points)
        set_cell_margins(
            cell,
            top=callout_theme.margin_top_dxa,
            bottom=callout_theme.margin_bottom_dxa,
            start=callout_theme.margin_start_dxa,
            end=callout_theme.margin_end_dxa,
        )

        row_pr = table.rows[0]._tr.get_or_add_trPr()
        row_pr.append(OxmlElement("w:cantSplit"))

        cell.text = ""
        label_p = cell.paragraphs[0]
        label_p.paragraph_format.space_after = Pt(callout_theme.label_space_after_pt)
        label_run = label_p.add_run((callout.label or style.label).upper())
        label_run.font.name = theme.typography.body_font
        label_run.bold = True
        label_run.font.size = Pt(theme.typography.callout_label_size_pt)
        label_run.font.color.rgb = rgb(style.code_colour)

        title_p = cell.add_paragraph()
        title_p.paragraph_format.space_after = Pt(callout_theme.title_space_after_pt)
        title_run = title_p.add_run(callout.title)
        title_run.font.name = theme.typography.display_font
        title_run.bold = True
        title_run.font.size = Pt(theme.typography.callout_title_size_pt)
        title_run.font.color.rgb = rgb(style.title_colour)
        if bookmark:
            self.add_bookmark(title_p, bookmark)

        for block in callout.body:
            if isinstance(block, BulletList):
                for item in block.items:
                    p = cell.add_paragraph(style="List Bullet")
                    text = item.text if isinstance(item, BulletItem) else str(item)
                    r = p.add_run(text)
                    r.font.name = theme.typography.body_font
                    r.font.size = Pt(theme.typography.callout_body_size_pt)
                    r.font.color.rgb = rgb(style.body_colour)
                continue

            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(callout_theme.body_space_after_pt)
            if block.inline_content:
                for item in block.inline_content:
                    if isinstance(item, CrossReference) and item.target_bookmark:
                        self.add_internal_hyperlink(p, item.render_label, item.target_bookmark)
                        continue
                    text = item.render_label if isinstance(item, CrossReference) else str(item)
                    r = p.add_run(text)
                    r.font.name = theme.typography.body_font
                    r.font.size = Pt(theme.typography.callout_body_size_pt)
                    r.font.color.rgb = rgb(style.accent if block.role == "bold" else style.body_colour)
                    r.bold = block.role == "bold"
            else:
                r = p.add_run(block.text)
                r.font.name = theme.typography.body_font
                r.font.size = Pt(theme.typography.callout_body_size_pt)
                r.font.color.rgb = rgb(style.body_colour)
                if block.role == "bold":
                    r.bold = True
                    r.font.color.rgb = rgb(style.accent)

        doc.add_paragraph().paragraph_format.space_after = Pt(callout_theme.trailing_space_after_pt)

    def flow_diagram(self, doc, flow: FlowDiagram) -> None:
        theme = self.theme.flow
        if flow.direction == "horizontal":
            text = f" {theme.arrow_glyph} ".join(node.label for node in flow.nodes)
            self.para(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, color=theme.node_colour, size=theme.node_size_pt, space_after=theme.arrow_space_after_pt)
            return
        for index, node in enumerate(flow.nodes):
            self.para(doc, node.label, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, color=theme.node_colour, size=theme.node_size_pt, space_after=theme.node_space_after_pt)
            if node.connector:
                self.para(doc, node.connector, align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, color=theme.connector_colour, size=theme.connector_size_pt, space_after=theme.connector_space_after_pt)
            if index < len(flow.nodes) - 1:
                self.para(doc, theme.arrow_glyph, align=WD_ALIGN_PARAGRAPH.CENTER, color=theme.connector_colour, size=theme.connector_size_pt, space_after=theme.arrow_space_after_pt)

    def part_title_page(self, doc, eyebrow: str, title: str, *, bookmark: str | None = None) -> None:
        theme = self.effective.volume_page
        p = self.para(doc, eyebrow, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=theme.eyebrow_size_pt, color=theme.eyebrow_colour, space_after=theme.eyebrow_space_after_pt)
        p.paragraph_format.space_before = Pt(theme.space_before_pt)
        p.paragraph_format.page_break_before = theme.page_break_before
        self.para(doc, title, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=theme.title_size_pt, color=theme.title_colour, space_after=theme.title_space_after_pt, bookmark=bookmark)
        if theme.page_break_after:
            doc.add_page_break()
