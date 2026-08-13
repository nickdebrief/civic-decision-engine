"""Format-specific output, link, accessibility, and equivalence validation."""

from __future__ import annotations

import importlib
import re
import subprocess
import zipfile
from dataclasses import dataclass, field
from html.parser import HTMLParser
from pathlib import Path
from xml.etree import ElementTree

from docx import Document

from model import Book, BulletList, Callout, Chapter, FlowDiagram, FrontMatter, Paragraph, Section, Volume
from renderers.pdf_renderer import discover_tool
from themes.base import EffectiveTheme
from validator import ValidationResult


WORD_NAMESPACE = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
VALID_IDENTIFIER = re.compile(r"^[A-Za-z][A-Za-z0-9_-]*$")


def normalize_text(value: str) -> str:
    return re.sub(r"\s+", " ", value).strip()


def source_text_blocks(book: Book) -> list[str]:
    """Return source-derived semantic text, excluding generated publication sections."""
    values: list[str] = []

    def visit(block) -> None:
        if isinstance(block, Section) and block.generated:
            return
        if isinstance(block, Volume):
            values.append(block.title)
        elif isinstance(block, FrontMatter):
            values.append(block.title)
        elif isinstance(block, Chapter):
            values.append(block.title)
        elif isinstance(block, Section):
            values.append(block.heading_text)
        elif isinstance(block, Callout):
            values.extend(part for part in (block.code or "", block.title) if part)
        elif isinstance(block, Paragraph):
            values.append(block.text)
        elif isinstance(block, BulletList):
            values.extend(item.text for item in block.items)
        elif isinstance(block, FlowDiagram):
            for node in block.nodes:
                values.append(node.label)
                if node.connector:
                    values.append(node.connector)
        for child in list(getattr(block, "blocks", [])) + list(getattr(block, "body", [])):
            visit(child)

    for root in book.blocks:
        visit(root)
    return [normalized for value in values if (normalized := normalize_text(value))]


def docx_text(path: Path) -> tuple[str, int, int]:
    document = Document(path)
    values = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                values.extend(paragraph.text for paragraph in cell.paragraphs)
    return normalize_text("\n".join(values)), len(document.paragraphs), len(document.tables)


@dataclass
class DocxAudit:
    paragraph_count: int = 0
    table_count: int = 0
    bookmark_count: int = 0
    hyperlink_count: int = 0
    duplicate_bookmarks: list[str] = field(default_factory=list)
    broken_links: list[str] = field(default_factory=list)


def audit_docx(path: Path) -> DocxAudit:
    _, paragraphs, tables = docx_text(path)
    with zipfile.ZipFile(path) as package:
        root = ElementTree.fromstring(package.read("word/document.xml"))
    names = [node.attrib.get(f"{{{WORD_NAMESPACE}}}name", "") for node in root.findall(f".//{{{WORD_NAMESPACE}}}bookmarkStart")]
    anchors = [node.attrib.get(f"{{{WORD_NAMESPACE}}}anchor", "") for node in root.findall(f".//{{{WORD_NAMESPACE}}}hyperlink") if node.attrib.get(f"{{{WORD_NAMESPACE}}}anchor")]
    duplicates = sorted({name for name in names if name and names.count(name) > 1})
    broken = sorted({anchor for anchor in anchors if anchor not in set(names)})
    return DocxAudit(paragraphs, tables, len(names), len(anchors), duplicates, broken)


def validate_docx_output(path: Path, book: Book) -> tuple[ValidationResult, DocxAudit]:
    result = ValidationResult()
    exists = path.is_file() and path.stat().st_size > 0
    result.add("DOCX generated", exists, str(path))
    if not exists:
        return result, DocxAudit()
    try:
        text, _, _ = docx_text(path)
        audit = audit_docx(path)
        properties = Document(path).core_properties
    except Exception as exc:
        result.add("DOCX reopened", False, str(exc))
        return result, DocxAudit()
    result.add("DOCX reopened", True, f"{audit.paragraph_count} paragraphs, {audit.table_count} tables")
    expected_tables = sum(isinstance(block, Callout) for block in _walk_model(book))
    result.add("DOCX structure", audit.paragraph_count > 0 and audit.table_count == expected_tables, f"{expected_tables} expected tables")
    result.add("DOCX metadata", bool(properties.title and properties.author and properties.language), properties.title or "missing")
    result.add("DOCX bookmarks unique", not audit.duplicate_bookmarks, ", ".join(audit.duplicate_bookmarks))
    result.add("DOCX internal links", not audit.broken_links, f"{audit.hyperlink_count} links, {len(audit.broken_links)} broken")
    result.add("DOCX unresolved markup", "[[REF:" not in text, "none" if "[[REF:" not in text else "found")
    return result, audit


def _walk_model(book: Book):
    stack = list(book.blocks)
    while stack:
        block = stack.pop(0)
        yield block
        stack[0:0] = list(getattr(block, "blocks", [])) + list(getattr(block, "body", []))


class HtmlAuditParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.language = ""
        self.ids: list[str] = []
        self.hrefs: list[str] = []
        self.asset_links: list[str] = []
        self.headings: list[tuple[int, str]] = []
        self.text_parts: list[str] = []
        self._heading_level: int | None = None
        self._heading_text: list[str] = []
        self._link_text_stack: list[list[str]] = []
        self.empty_links = 0

    def handle_starttag(self, tag: str, attrs) -> None:
        attributes = dict(attrs)
        if tag == "html":
            self.language = attributes.get("lang", "")
        identifier = attributes.get("id")
        if identifier is not None:
            self.ids.append(identifier)
        if tag == "a":
            href = attributes.get("href", "")
            self.hrefs.append(href)
            self._link_text_stack.append([])
        if tag in {"img", "link", "script"}:
            target = attributes.get("src") or attributes.get("href")
            if target:
                self.asset_links.append(target)
        if tag in {"h1", "h2", "h3", "h4", "h5", "h6"}:
            self._heading_level = int(tag[1])
            self._heading_text = []

    def handle_endtag(self, tag: str) -> None:
        if tag == "a" and self._link_text_stack:
            if not normalize_text("".join(self._link_text_stack.pop())):
                self.empty_links += 1
        if self._heading_level is not None and tag == f"h{self._heading_level}":
            self.headings.append((self._heading_level, normalize_text("".join(self._heading_text))))
            self._heading_level = None
            self._heading_text = []

    def handle_data(self, data: str) -> None:
        self.text_parts.append(data)
        if self._heading_level is not None:
            self._heading_text.append(data)
        if self._link_text_stack:
            self._link_text_stack[-1].append(data)

    @property
    def text(self) -> str:
        return normalize_text(" ".join(self.text_parts))


@dataclass
class HtmlAudit:
    anchor_count: int = 0
    internal_link_count: int = 0
    duplicate_ids: list[str] = field(default_factory=list)
    broken_links: list[str] = field(default_factory=list)
    invalid_ids: list[str] = field(default_factory=list)
    empty_headings: int = 0
    heading_gaps: list[str] = field(default_factory=list)
    empty_links: int = 0
    broken_assets: list[str] = field(default_factory=list)
    language: str = ""
    text: str = ""


def parse_html_output(path: Path) -> tuple[HtmlAuditParser, Path]:
    html_path = path / "index.html" if path.is_dir() else path
    parser = HtmlAuditParser()
    parser.feed(html_path.read_text(encoding="utf-8"))
    return parser, html_path


def audit_html(path: Path) -> HtmlAudit:
    parser, html_path = parse_html_output(path)
    duplicates = sorted({identifier for identifier in parser.ids if parser.ids.count(identifier) > 1})
    internal = [href[1:] for href in parser.hrefs if href.startswith("#")]
    broken = sorted({target for target in internal if target not in set(parser.ids)})
    invalid = sorted({identifier for identifier in parser.ids if not VALID_IDENTIFIER.fullmatch(identifier)})
    empty_headings = sum(not text for _, text in parser.headings)
    gaps = []
    previous = None
    for level, text in parser.headings:
        if previous is not None and level > previous + 1:
            gaps.append(text)
        previous = level
    broken_assets = []
    for target in parser.asset_links:
        if target.startswith(("data:", "http://", "https://", "#")):
            continue
        if not (html_path.parent / target).is_file():
            broken_assets.append(target)
    return HtmlAudit(
        anchor_count=len(parser.ids),
        internal_link_count=len(internal),
        duplicate_ids=duplicates,
        broken_links=broken,
        invalid_ids=invalid,
        empty_headings=empty_headings,
        heading_gaps=gaps,
        empty_links=parser.empty_links,
        broken_assets=broken_assets,
        language=parser.language,
        text=parser.text,
    )


def validate_html_output(path: Path, language: str) -> tuple[ValidationResult, HtmlAudit]:
    result = ValidationResult()
    html_path = path / "index.html" if path.is_dir() else path
    exists = html_path.is_file() and html_path.stat().st_size > 0
    result.add("HTML generated", exists, str(html_path))
    if not exists:
        return result, HtmlAudit()
    try:
        audit = audit_html(path)
    except Exception as exc:
        result.add("HTML parsed", False, str(exc))
        return result, HtmlAudit()
    result.add("HTML parsed", True, f"{audit.anchor_count} anchors, {audit.internal_link_count} internal links")
    result.add("HTML language", audit.language == language, audit.language or "missing")
    result.add("HTML identifiers", not audit.duplicate_ids and not audit.invalid_ids, f"{len(audit.duplicate_ids)} duplicate, {len(audit.invalid_ids)} invalid")
    result.add("HTML internal links", not audit.broken_links and audit.empty_links == 0, f"{len(audit.broken_links)} broken, {audit.empty_links} empty")
    result.add("HTML assets", not audit.broken_assets, ", ".join(audit.broken_assets))
    result.add("HTML heading hierarchy", audit.empty_headings == 0 and not audit.heading_gaps, f"{audit.empty_headings} empty, {len(audit.heading_gaps)} gaps")
    result.add("HTML unresolved markup", "[[REF:" not in audit.text, "none" if "[[REF:" not in audit.text else "found")
    return result, audit


@dataclass(frozen=True)
class PdfTextExtraction:
    text: str = ""
    raw_text: str = ""
    status: str = "unavailable"
    backend: str = ""
    reason: str = "No supported PDF text extraction backend was available."
    attempts: tuple[tuple[str, str], ...] = ()
    page_count: int = 0
    width_points: float = 0.0
    height_points: float = 0.0
    title: str = ""

    @property
    def available(self) -> bool:
        return self.status == "available"


def _pypdf_metadata(reader, pages) -> tuple[int, float, float, str]:
    width = float(pages[0].mediabox.width) if pages else 0.0
    height = float(pages[0].mediabox.height) if pages else 0.0
    metadata = reader.metadata or {}
    title = getattr(metadata, "title", None)
    if not title and hasattr(metadata, "get"):
        title = metadata.get("/Title", "")
    return len(pages), width, height, str(title or "")


def extract_pdf_text_result(path: Path) -> PdfTextExtraction:
    """Extract PDF text with deterministic, ordered optional backends."""

    attempts: list[tuple[str, str]] = []
    tool = discover_tool("pdftotext")
    if tool is not None:
        completed = subprocess.run(
            [str(tool), "-layout", str(path), "-"],
            check=False,
            capture_output=True,
            text=True,
            timeout=120,
        )
        if completed.returncode != 0:
            raise RuntimeError(f"PDF text extraction failed: {completed.stderr.strip()}")
        attempts.append(("pdftotext", "available"))
        return PdfTextExtraction(
            text=normalize_text(completed.stdout),
            raw_text=completed.stdout,
            status="available",
            backend="pdftotext",
            reason="",
            attempts=tuple(attempts),
        )
    attempts.append(("pdftotext", "not found"))

    try:
        pypdf = importlib.import_module("pypdf")
    except ImportError:
        attempts.append(("pypdf", "module unavailable"))
    else:
        try:
            reader = pypdf.PdfReader(str(path))
            pages = list(reader.pages)
            text = normalize_text("\n".join(page.extract_text() or "" for page in pages))
            page_count, width, height, title = _pypdf_metadata(reader, pages)
        except Exception as exc:
            attempts.append(("pypdf", f"extraction failed: {type(exc).__name__}"))
        else:
            attempts.append(("pypdf", "available"))
            return PdfTextExtraction(
                text=text,
                raw_text="\n".join(page.extract_text() or "" for page in pages),
                status="available",
                backend="pypdf",
                reason="",
                attempts=tuple(attempts),
                page_count=page_count,
                width_points=width,
                height_points=height,
                title=title,
            )

    try:
        pdfminer = importlib.import_module("pdfminer.high_level")
    except ImportError:
        attempts.append(("pdfminer.six", "module unavailable"))
    else:
        try:
            text = normalize_text(str(pdfminer.extract_text(str(path)) or ""))
        except Exception as exc:
            attempts.append(("pdfminer.six", f"extraction failed: {type(exc).__name__}"))
        else:
            attempts.append(("pdfminer.six", "available"))
            return PdfTextExtraction(
                text=text,
                raw_text=str(pdfminer.extract_text(str(path)) or ""),
                status="available",
                backend="pdfminer.six",
                reason="",
                attempts=tuple(attempts),
            )

    return PdfTextExtraction(attempts=tuple(attempts))


def extract_pdf_text(path: Path) -> str:
    """Compatibility API returning text only; unavailable extraction is empty."""

    return extract_pdf_text_result(path).text


PDF_RUNNING_HEADER = re.compile(
    r"Structured\s*·\s*Traceable\s*·\s*Governed\s*·\s*\d+\s+EVIDENCE-LED GOVERNANCE"
)


def _pdf_equivalence_pattern(block: str) -> re.Pattern[str]:
    """Build a strict pattern for PDF line-wrap representation artifacts."""
    escaped = re.escape(block)
    escaped = escaped.replace(r"\ ", r"\s+")
    escaped = re.sub(
        r"\\-(?=[A-Za-z])",
        lambda _match: r"(?:-\s*|[\r\n\f]\s*)",
        escaped,
    )
    return re.compile(escaped, re.DOTALL)


def pdf_block_matches(block: str, extraction: PdfTextExtraction) -> bool:
    """Match source text while accounting only for proven PDF representation artifacts."""
    if block in extraction.text:
        return True
    raw = extraction.raw_text or extraction.text
    raw = PDF_RUNNING_HEADER.sub(" ", raw)
    return _pdf_equivalence_pattern(block).search(raw) is not None


@dataclass
class PdfAudit:
    page_count: int = 0
    width_points: float = 0.0
    height_points: float = 0.0
    title: str = ""
    text: str = ""
    inspection_available: bool = False
    validation_status: str = "unavailable"
    text_backend: str = ""
    validation_reason: str = "No supported PDF text extraction backend was available."
    backend_attempts: tuple[tuple[str, str], ...] = ()


def audit_pdf(path: Path) -> PdfAudit:
    extraction = extract_pdf_text_result(path)
    values: dict[str, str] = {}
    pdfinfo = discover_tool("pdfinfo")
    if pdfinfo is not None:
        completed = subprocess.run(
            [str(pdfinfo), str(path)],
            check=False,
            capture_output=True,
            text=True,
            timeout=120,
        )
        if completed.returncode != 0:
            raise RuntimeError(f"PDF inspection failed: {completed.stderr.strip()}")
        for line in completed.stdout.splitlines():
            if ":" in line:
                key, value = line.split(":", 1)
                values[key.strip()] = value.strip()
    size = re.search(r"([0-9.]+)\s+x\s+([0-9.]+)\s+pts", values.get("Page size", ""))
    page_count = int(values.get("Pages", "0")) if values else extraction.page_count
    width = float(size.group(1)) if size else extraction.width_points
    height = float(size.group(2)) if size else extraction.height_points
    title = values.get("Title", "") if values else extraction.title
    return PdfAudit(
        page_count=page_count,
        width_points=width,
        height_points=height,
        title=title,
        text=extraction.text,
        inspection_available=bool(values) or extraction.page_count > 0,
        validation_status=extraction.status,
        text_backend=extraction.backend,
        validation_reason=extraction.reason,
        backend_attempts=extraction.attempts,
    )


def validate_pdf_output(path: Path, book: Book, effective: EffectiveTheme) -> tuple[ValidationResult, PdfAudit]:
    result = ValidationResult()
    exists = path.is_file() and path.stat().st_size > 0
    result.add("PDF generated", exists, str(path))
    if not exists:
        return result, PdfAudit()
    try:
        audit = audit_pdf(path)
    except Exception as exc:
        result.add("PDF opened", False, str(exc))
        return result, PdfAudit()
    if audit.inspection_available:
        expected_width = effective.page.width_inches * 72
        expected_height = effective.page.height_inches * 72
        dimensions_match = abs(audit.width_points - expected_width) <= 2 and abs(audit.height_points - expected_height) <= 2
        result.add("PDF opened", audit.page_count > 0, f"{audit.page_count} pages")
        result.add("PDF metadata", bool(audit.title), audit.title or "missing title")
        result.add("PDF page profile", dimensions_match, f"{audit.width_points:g} x {audit.height_points:g} pt")
    else:
        result.add("PDF structural validation", True, "Skipped: no supported inspection backend was available")
    if audit.validation_status == "available":
        result.add("PDF expected text", book.title in audit.text and book.author in audit.text, book.title)
    else:
        result.add("PDF expected text", True, f"Skipped: {audit.validation_reason}")
    return result, audit


@dataclass
class EquivalenceAudit:
    block_count: int = 0
    missing_docx: list[str] = field(default_factory=list)
    missing_html: list[str] = field(default_factory=list)
    missing_pdf: list[str] = field(default_factory=list)
    pdf_status: str = "not requested"
    pdf_backend: str = ""
    pdf_reason: str = ""
    pdf_attempts: tuple[tuple[str, str], ...] = ()


def validate_cross_format_equivalence(
    book: Book,
    *,
    docx_path: Path,
    html_path: Path | None = None,
    pdf_path: Path | None = None,
) -> tuple[ValidationResult, EquivalenceAudit]:
    result = ValidationResult()
    blocks = source_text_blocks(book)
    docx_value, _, _ = docx_text(docx_path)
    html_value = audit_html(html_path).text if html_path is not None else ""
    pdf_extraction = extract_pdf_text_result(pdf_path) if pdf_path is not None else None
    pdf_value = pdf_extraction.text if pdf_extraction is not None else ""
    missing_docx = [block for block in blocks if block not in docx_value]
    missing_html = [block for block in blocks if html_path is not None and block not in html_value]
    missing_pdf = [
        block
        for block in blocks
        if pdf_extraction is not None
        and pdf_extraction.available
        and not pdf_block_matches(block, pdf_extraction)
    ]
    result.add("DOCX source equivalence", not missing_docx, f"{len(blocks) - len(missing_docx)}/{len(blocks)} blocks")
    if html_path is not None:
        result.add("HTML source equivalence", not missing_html, f"{len(blocks) - len(missing_html)}/{len(blocks)} blocks")
    if pdf_extraction is not None:
        if pdf_extraction.available:
            result.add("PDF source equivalence", not missing_pdf, f"{len(blocks) - len(missing_pdf)}/{len(blocks)} blocks")
        else:
            result.add("PDF source equivalence", True, f"Skipped: {pdf_extraction.reason}")
    return result, EquivalenceAudit(
        len(blocks),
        missing_docx,
        missing_html,
        missing_pdf,
        pdf_status=(pdf_extraction.status if pdf_extraction is not None else "not requested"),
        pdf_backend=(pdf_extraction.backend if pdf_extraction is not None else ""),
        pdf_reason=(pdf_extraction.reason if pdf_extraction is not None else ""),
        pdf_attempts=(pdf_extraction.attempts if pdf_extraction is not None else ()),
    )
