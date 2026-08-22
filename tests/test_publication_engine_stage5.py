from __future__ import annotations

import json
import subprocess
import sys
import tempfile
import unittest
import zipfile
from pathlib import Path
from unittest.mock import patch

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


PIPELINE_DIR = Path(__file__).resolve().parents[1] / "scripts" / "evidence_led_governance_pipeline"
REPOSITORY_DIR = PIPELINE_DIR.parents[1]
sys.path.insert(0, str(PIPELINE_DIR))

from build import (  # noqa: E402
    AUTHOR,
    BASENAME,
    RUNNING_TITLE,
    SUBTITLE,
    TAGLINE,
    TITLE,
    PublicationBuildError,
    _assert_promotable,
    build_publication,
)
from manifest import load_manifest  # noqa: E402
from model import (  # noqa: E402
    Book,
    BulletItem,
    BulletList,
    Callout,
    Chapter,
    FlowDiagram,
    FlowNode,
    GovernancePrinciple,
    HtmlOutputConfig,
    Manifest,
    OutputConfig,
    PackageOutputConfig,
    Paragraph,
    PublicationConfig,
    Section,
    Volume,
)
from output_validation import (  # noqa: E402
    audit_docx,
    audit_html,
    validate_cross_format_equivalence,
    validate_docx_output,
    validate_html_output,
    validate_pdf_output,
    source_text_blocks,
)
from packaging import create_package, sha256_file  # noqa: E402
from parser import Parser  # noqa: E402
from publication import enrich_publication  # noqa: E402
from renderers.docx_renderer import DocxRenderer  # noqa: E402
from renderers.html_renderer import HtmlRenderer, generate_css  # noqa: E402
from renderers.pdf_renderer import PdfRenderer  # noqa: E402
from theme_resolution import resolve_theme  # noqa: E402
from validator import validate_book  # noqa: E402


def write(path: Path, text: str) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text, encoding="utf-8")
    return path


def manifest_text(*, formats: str = '"docx", "pdf", "html"', theme: str = "handbook", html_extra: str = "") -> str:
    return f'''schema_version = 1
sources = ["chapters/chapter.txt"]

[publication]
title = "Test Publication"
subtitle = "Evidence test"
author = "Test Author"
language = "en"
edition = "Test"
theme = "{theme}"

[output]
basename = "Test_Publication"
directory = "output"
formats = [{formats}]
profile = "digital"

[output.html]
single_file = true
include_navigation = true
include_semantic_index = true
embed_css = true
{html_extra}

[output.pdf]
source = "docx"
require_render = true
preserve_bookmarks = true

[output.package]
enabled = true
include_checksums = true
include_build_report = true

[layout]
page_profile = "letter"
'''


def load_fixture_manifest(root: Path, **kwargs) -> Manifest:
    write(root / "chapters/chapter.txt", "# Chapter 1 — Test\n\n## 1.1 Evidence\n\nBody evidence.\n")
    manifest_path = write(root / "book.toml", manifest_text(**kwargs))
    return load_manifest(manifest_path, chapters_dir=root / "chapters")


def minimal_book() -> Book:
    callout = GovernancePrinciple(
        label="Governance Principle",
        title="GP-1 — Rights Require Governance",
        code="GP-1",
        callout_type="Governance Principle",
        body=[Paragraph(text="Callout evidence.")],
    )
    section = Section(number="1.1", title="Evidence", blocks=[Paragraph(text="Body evidence."), BulletList(items=[BulletItem(text="Evidence item.")]), callout])
    chapter = Chapter(number=1, title="Test Chapter", sections=[section], blocks=[section])
    volume = Volume(number="I", title="Test Volume", chapters=[chapter], blocks=[chapter])
    book = Book(
        title="Test Publication",
        subtitle="Evidence test",
        author="Test Author",
        running_title="TEST PUBLICATION",
        tagline="Structured · Traceable · Governed",
        version="1.0",
        metadata={"language": "en", "edition": "Test"},
        volumes=[volume],
        blocks=[volume],
    )
    enrich_publication(book)
    return book


def effective_theme(name: str = "handbook"):
    return resolve_theme(Manifest(publication=PublicationConfig(theme=name)), assets_dir=PIPELINE_DIR / "assets").effective_theme


class PublicationEngineStage5Tests(unittest.TestCase):
    def test_manifest_accepts_multiple_output_formats(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            manifest = load_fixture_manifest(Path(temp))
        self.assertEqual(manifest.output.formats, ("docx", "pdf", "html"))

    def test_unknown_output_format_fails(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            with self.assertRaisesRegex(ValueError, "unsupported output format"):
                load_fixture_manifest(root, formats='"epub"')

    def test_duplicate_output_format_fails(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            with self.assertRaisesRegex(ValueError, "duplicates"):
                load_fixture_manifest(Path(temp), formats='"docx", "docx"')

    def test_manifest_loads_html_pdf_and_package_profiles(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            manifest = load_fixture_manifest(Path(temp))
        self.assertTrue(manifest.output.html.single_file)
        self.assertEqual(manifest.output.pdf.source, "docx")
        self.assertTrue(manifest.output.package.enabled)

    def test_unknown_pdf_source_fails(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            write(root / "chapters/chapter.txt", "# Chapter 1 — Test")
            path = write(root / "book.toml", manifest_text().replace('source = "docx"', 'source = "html"'))
            with self.assertRaisesRegex(ValueError, "must be 'docx'"):
                load_manifest(path, chapters_dir=root / "chapters")

    def test_single_file_html_requires_embedded_css(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            text = manifest_text().replace("embed_css = true", "embed_css = false")
            write(root / "chapters/chapter.txt", "# Chapter 1 — Test")
            path = write(root / "book.toml", text)
            with self.assertRaisesRegex(ValueError, "requires"):
                load_manifest(path, chapters_dir=root / "chapters")

    def test_html_renderer_produces_semantic_structure_and_language(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            out = Path(temp) / "book.html"
            HtmlRenderer(effective_theme(), HtmlOutputConfig()).render(minimal_book(), out)
            source = out.read_text(encoding="utf-8")
        self.assertIn('<html lang="en">', source)
        for tag in ("<header", "<nav", "<main", "<section", "<article", "<aside", "<footer"):
            self.assertIn(tag, source)

    def test_html_contains_accessible_skip_link_and_navigation_labels(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            out = Path(temp) / "book.html"
            HtmlRenderer(effective_theme(), HtmlOutputConfig()).render(minimal_book(), out)
            source = out.read_text(encoding="utf-8")
        self.assertIn('class="skip-link"', source)
        self.assertIn('aria-label="Publication navigation"', source)
        self.assertIn('aria-labelledby=', source)

    def test_html_ids_and_css_are_deterministic(self) -> None:
        book = minimal_book()
        theme = effective_theme()
        with tempfile.TemporaryDirectory() as temp:
            first = Path(temp) / "first.html"
            second = Path(temp) / "second.html"
            HtmlRenderer(theme, HtmlOutputConfig()).render(book, first)
            HtmlRenderer(theme, HtmlOutputConfig()).render(book, second)
            self.assertEqual(first.read_bytes(), second.read_bytes())
        self.assertEqual(generate_css(theme), generate_css(theme))

    def test_html_links_resolve_and_callout_type_is_preserved(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            out = Path(temp) / "book.html"
            HtmlRenderer(effective_theme(), HtmlOutputConfig()).render(minimal_book(), out)
            audit = audit_html(out)
            source = out.read_text(encoding="utf-8")
        self.assertFalse(audit.broken_links)
        self.assertIn("callout-governance-principle", source)

    def test_html_accessibility_audit_has_no_empty_headings_or_gaps(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            out = Path(temp) / "book.html"
            HtmlRenderer(effective_theme(), HtmlOutputConfig()).render(minimal_book(), out)
            audit = audit_html(out)
        self.assertEqual(audit.empty_headings, 0)
        self.assertFalse(audit.heading_gaps)

    def test_html_semantic_index_and_no_raw_reference_markup(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            out = Path(temp) / "book.html"
            HtmlRenderer(effective_theme(), HtmlOutputConfig()).render(minimal_book(), out)
            source = out.read_text(encoding="utf-8")
        self.assertIn("Semantic Index", source)
        self.assertNotIn("[[REF:", source)

    def test_single_file_and_directory_html_builds_succeed(self) -> None:
        book = minimal_book()
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            HtmlRenderer(effective_theme(), HtmlOutputConfig()).render(book, root / "single.html")
            directory = root / "site"
            HtmlRenderer(effective_theme(), HtmlOutputConfig(single_file=False, embed_css=False)).render(book, directory)
            self.assertTrue((directory / "index.html").is_file())
            self.assertTrue((directory / "styles.css").is_file())

    def test_directory_html_output_validation_succeeds(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            directory = Path(temp) / "site"
            HtmlRenderer(effective_theme(), HtmlOutputConfig(single_file=False, embed_css=False)).render(minimal_book(), directory)
            result, audit = validate_html_output(directory, "en")
        self.assertTrue(result.ok)
        self.assertGreater(audit.anchor_count, 0)

    def test_html_validation_catches_broken_anchor_and_duplicate_id(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            path = write(Path(temp) / "bad.html", '<!doctype html><html lang="en"><body><h1 id="same">One</h1><p id="same"><a href="#missing">Missing</a></p></body></html>')
            result, audit = validate_html_output(path, "en")
        self.assertFalse(result.ok)
        self.assertEqual(audit.duplicate_ids, ["same"])
        self.assertEqual(audit.broken_links, ["missing"])

    def test_pdf_renderer_detects_missing_libreoffice(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            source = Path(temp) / "source.docx"
            Document().save(source)
            with self.assertRaisesRegex(RuntimeError, "LibreOffice"):
                PdfRenderer(Path(temp) / "missing-soffice").render(source, Path(temp) / "out.pdf")

    def test_pdf_renderer_accepts_libreoffice_alias_when_soffice_is_absent(self) -> None:
        import renderers.pdf_renderer as pdf_renderer

        with tempfile.TemporaryDirectory() as temp:
            executable = Path(temp) / "libreoffice"
            executable.write_text("#!/bin/sh\nprintf 'LibreOffice synthetic\\n'\n", encoding="utf-8")
            executable.chmod(0o700)
            with patch.object(pdf_renderer, "discover_tool", side_effect=lambda name: executable if name == "libreoffice" else None):
                renderer = PdfRenderer()
                self.assertEqual(renderer.soffice_path, executable)
                self.assertTrue(renderer.available)

    def test_pdf_renderer_discovery_prefers_usable_entry_point_and_falls_back(self) -> None:
        import renderers.pdf_renderer as pdf_renderer

        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            soffice = root / "soffice"
            libreoffice = root / "libreoffice"
            for path in (soffice, libreoffice):
                path.write_text("#!/bin/sh\nprintf 'LibreOffice synthetic\\n'\n", encoding="utf-8")
                path.chmod(0o700)

            def discover(name: str):
                return {"soffice": soffice, "libreoffice": libreoffice}.get(name)

            with patch.object(pdf_renderer, "discover_tool", side_effect=discover):
                self.assertEqual(PdfRenderer().soffice_path, soffice)

            libreoffice.unlink()
            with patch.object(pdf_renderer, "discover_tool", side_effect=discover):
                self.assertEqual(PdfRenderer().soffice_path, soffice)

            libreoffice.write_text("#!/bin/sh\nprintf 'LibreOffice synthetic\\n'\n", encoding="utf-8")
            libreoffice.chmod(0o700)
            soffice.chmod(0o600)
            with patch.object(pdf_renderer, "discover_tool", side_effect=discover):
                self.assertEqual(PdfRenderer().soffice_path, libreoffice)

            soffice.chmod(0o700)
            with patch.object(pdf_renderer, "discover_tool", side_effect=discover), patch.object(
                pdf_renderer.subprocess, "run", side_effect=lambda command, **_: subprocess.CompletedProcess(command, 1 if command[0] == str(soffice) else 0, stdout="", stderr="bad")
            ):
                self.assertEqual(PdfRenderer().soffice_path, libreoffice)

    def test_pdf_renderer_accepts_symlink_and_shell_wrapper_when_version_succeeds(self) -> None:
        import renderers.pdf_renderer as pdf_renderer

        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            wrapper = root / "libreoffice-wrapper"
            wrapper.write_text("#!/bin/sh\nprintf 'LibreOffice synthetic\\n'\n", encoding="utf-8")
            wrapper.chmod(0o700)
            symlink = root / "soffice"
            symlink.symlink_to(wrapper)
            with patch.object(pdf_renderer, "discover_tool", side_effect=lambda name: symlink if name == "soffice" else None):
                renderer = PdfRenderer()
            self.assertEqual(renderer.soffice_path, symlink)
            self.assertIn("LibreOffice synthetic", renderer.version())

    def test_pdf_renderer_rejects_unusable_entry_points_and_reports_selected_identity(self) -> None:
        import renderers.pdf_renderer as pdf_renderer

        with tempfile.TemporaryDirectory() as temp:
            unusable = Path(temp) / "soffice"
            unusable.write_text("#!/bin/sh\nexit 1\n", encoding="utf-8")
            unusable.chmod(0o700)
            with patch.object(pdf_renderer, "discover_tool", side_effect=lambda name: unusable if name in {"soffice", "libreoffice"} else None):
                renderer = PdfRenderer()
            self.assertIsNone(renderer.soffice_path)
            self.assertFalse(renderer.available)

    def test_pdf_renderer_and_validation_succeed_when_available(self) -> None:
        renderer = PdfRenderer()
        if not renderer.available:
            self.skipTest("LibreOffice unavailable")
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            book = minimal_book()
            docx = root / "book.docx"
            pdf = root / "book.pdf"
            DocxRenderer(effective_theme()).render(book, docx)
            renderer.render(docx, pdf)
            result, audit = validate_pdf_output(pdf, book, effective_theme())
        self.assertTrue(result.ok)
        self.assertGreater(audit.page_count, 0)
        self.assertAlmostEqual(audit.width_points, 612, delta=2)

    def test_pdf_validation_rejects_empty_output(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            empty = write(Path(temp) / "empty.pdf", "")
            result, audit = validate_pdf_output(empty, minimal_book(), effective_theme())
        self.assertFalse(result.ok)
        self.assertEqual(audit.page_count, 0)

    def test_docx_output_and_link_validation_passes(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            path = Path(temp) / "book.docx"
            book = minimal_book()
            DocxRenderer(effective_theme()).render(book, path)
            result, audit = validate_docx_output(path, book)
        self.assertTrue(result.ok)
        self.assertGreater(audit.bookmark_count, 0)
        self.assertGreater(audit.hyperlink_count, 0)

    def test_duplicate_docx_bookmark_fails_validation(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            path = Path(temp) / "duplicate.docx"
            document = Document()
            for index in (1, 2):
                paragraph = document.add_paragraph(f"Bookmark {index}")
                start = OxmlElement("w:bookmarkStart")
                start.set(qn("w:id"), str(index))
                start.set(qn("w:name"), "duplicate")
                paragraph._p.insert(0, start)
            document.save(path)
            audit = audit_docx(path)
        self.assertEqual(audit.duplicate_bookmarks, ["duplicate"])

    def test_docx_audit_detects_broken_internal_hyperlink(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            path = Path(temp) / "broken.docx"
            document = Document()
            paragraph = document.add_paragraph()
            hyperlink = OxmlElement("w:hyperlink")
            hyperlink.set(qn("w:anchor"), "missing_target")
            paragraph._p.append(hyperlink)
            document.save(path)
            audit = audit_docx(path)
        self.assertEqual(audit.broken_links, ["missing_target"])

    def test_cross_format_source_equivalence_and_missing_html_detection(self) -> None:
        renderer = PdfRenderer()
        if not renderer.available:
            self.skipTest("LibreOffice unavailable")
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            book = minimal_book()
            docx = root / "book.docx"
            html = root / "book.html"
            pdf = root / "book.pdf"
            DocxRenderer(effective_theme()).render(book, docx)
            HtmlRenderer(effective_theme(), HtmlOutputConfig()).render(book, html)
            renderer.render(docx, pdf)
            result, audit = validate_cross_format_equivalence(book, docx_path=docx, html_path=html, pdf_path=pdf)
            self.assertTrue(result.ok, (audit.missing_docx, audit.missing_html, audit.missing_pdf))
            html.write_text(html.read_text(encoding="utf-8").replace("Body evidence.", ""), encoding="utf-8")
            broken, broken_audit = validate_cross_format_equivalence(book, docx_path=docx, html_path=html)
        self.assertFalse(broken.ok)
        self.assertIn("Body evidence.", broken_audit.missing_html)

    def test_empty_list_fails_stronger_model_validation(self) -> None:
        book = minimal_book()
        chapter = book.volumes[0].chapters[0]
        chapter.sections[0].blocks.append(BulletList())
        self.assertFalse(validate_book(book).ok)

    def test_empty_callout_and_short_flow_fail_model_validation(self) -> None:
        book = minimal_book()
        section = book.volumes[0].chapters[0].sections[0]
        section.blocks.append(Callout(code="GP-99", title="Empty", body=[]))
        section.blocks.append(FlowDiagram(nodes=[FlowNode(label="Only")]))
        checks = {name: ok for name, ok, _ in validate_book(book).checks}
        self.assertFalse(checks["Empty Coded Object Bodies"])
        self.assertFalse(checks["Flow Diagram Shape"])

    def test_chapter_without_sections_fails_model_validation(self) -> None:
        book = minimal_book()
        book.volumes[0].chapters[0].sections.clear()
        checks = {name: ok for name, ok, _ in validate_book(book).checks}
        self.assertFalse(checks["Chapter Section Membership"])

    def test_package_manifest_checksums_and_report_are_generated(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            source = write(root / "chapter.txt", "Evidence")
            artifact = write(root / "book.html", "<html>Evidence</html>")
            manifest_path = write(root / "book.toml", "schema_version = 1")
            manifest = Manifest(path=manifest_path, source_files=[source], output=OutputConfig(package=PackageOutputConfig(enabled=True)))
            result = create_package(
                package_dir=root,
                stem="book_v1.0",
                repository=REPOSITORY_DIR,
                book=minimal_book(),
                manifest=manifest,
                artifacts={"html": artifact},
                build_report_text="PASS\n",
                renderer_versions={"html": "test"},
                validation_status="PASS",
            )
            data = json.loads(result.manifest.read_text(encoding="utf-8"))
            checksum_lines = result.checksums.read_text(encoding="utf-8")
            expected_checksum = sha256_file(artifact)
            expected_size = artifact.stat().st_size
        self.assertEqual(data["outputs"][0]["sha256"], expected_checksum)
        self.assertEqual(data["outputs"][0]["size_bytes"], expected_size)
        self.assertEqual(result.checksums.name, "book_v1.0_checksums.txt")
        self.assertIn("book_v1.0_build_report.txt", checksum_lines)

    def test_sha256_checksum_matches_known_value(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            path = write(Path(temp) / "value.txt", "evidence")
            self.assertEqual(sha256_file(path), "ee8250fb76e094b34b471f13a73dbbe51d1ae142e9df59d7c0d31ec20f0a0a8e")

    def test_source_text_blocks_exclude_generated_front_matter(self) -> None:
        blocks = source_text_blocks(minimal_book())
        self.assertIn("Body evidence.", blocks)
        self.assertNotIn("Semantic Index", blocks)

    def test_partial_promotion_refuses_existing_output(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            existing = write(Path(temp) / "book.docx", "authoritative")
            with self.assertRaises(PublicationBuildError):
                _assert_promotable([existing])
            self.assertEqual(existing.read_text(encoding="utf-8"), "authoritative")

    def test_legacy_docx_only_manifest_defaults_remain(self) -> None:
        manifest = Manifest()
        self.assertEqual(manifest.output.formats, ("docx",))
        self.assertFalse(manifest.output.package.enabled)

    def test_handbook_cde_and_cref_html_smoke_builds(self) -> None:
        book = minimal_book()
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            outputs = []
            for name in ("handbook", "cde", "cref"):
                out = root / f"{name}.html"
                HtmlRenderer(effective_theme(name), HtmlOutputConfig()).render(book, out)
                outputs.append(out.read_text(encoding="utf-8"))
        self.assertEqual(len(set(outputs)), 3)

    def test_theme_css_differs_without_changing_semantic_markup(self) -> None:
        handbook = generate_css(effective_theme("handbook"))
        cde = generate_css(effective_theme("cde"))
        cref = generate_css(effective_theme("cref"))
        self.assertEqual(len({handbook, cde, cref}), 3)
        self.assertIn(".callout-governance-principle", handbook)

    def test_existing_chapter_sources_are_not_modified(self) -> None:
        sources = sorted((PIPELINE_DIR / "chapters").glob("*.txt"))
        before = {path: path.read_bytes() for path in sources}
        Parser().parse_files(sources, title=TITLE, subtitle=SUBTITLE, author=AUTHOR, running_title=RUNNING_TITLE, tagline=TAGLINE, version="test")
        self.assertEqual(before, {path: path.read_bytes() for path in sources})

    def test_requested_pdf_failure_returns_rendering_exit_category_and_no_outputs(self) -> None:
        source_paths = sorted((PIPELINE_DIR / "chapters").glob("*.txt"))
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            source_lines = ",\n".join(f'  "{path}"' for path in source_paths)
            manifest_path = write(root / "book.toml", f'''schema_version = 1
sources = [
{source_lines}
]
[publication]
theme = "handbook"
[version]
mode = "fixed"
start = "1.0"
[output]
basename = "Failure_Test"
directory = "output"
formats = ["docx", "pdf"]
profile = "digital"
[output.pdf]
source = "docx"
require_render = true
[layout]
page_profile = "letter"
''')
            with patch("build.PdfRenderer", return_value=PdfRenderer(root / "missing")):
                with self.assertRaises(PublicationBuildError) as raised:
                    build_publication(
                        chapters_dir=PIPELINE_DIR / "chapters", output_dir=root / "fallback", basename=BASENAME,
                        title=TITLE, subtitle=SUBTITLE, author=AUTHOR, running_title=RUNNING_TITLE,
                        tagline=TAGLINE, manifest_path=manifest_path,
                    )
            self.assertEqual(raised.exception.exit_code, 3)
            self.assertFalse(any((root / "output").glob("Failure_Test*")))


if __name__ == "__main__":
    unittest.main()
